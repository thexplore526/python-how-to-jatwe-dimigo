from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION
from docx.oxml.ns import qn
from datetime import datetime
import os

with open(os.path.dirname(os.path.realpath(__file__)) + '/profile.txt', 'r', encoding='utf-8') as f:
    profiles = f.readlines()
profiles = list(map(lambda x: x.replace('\n',''), profiles))

name = profiles[0]
parent_name = profiles[1]
department = profiles[2]
birthday = profiles[3]
number = profiles[4]
phone = profiles[5]
address = profiles[6]
reason = profiles[7]

doc = Document()
header = doc.add_paragraph()
title = header.add_run('자퇴서')
title.font.size = shared.Pt(33)
title.font.name = '맑은 고딕'
title._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕') # 제목 글씨체
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows = 7, cols = 4)
table.style = doc.styles['Table Grid']

for i in range(5):
    table.rows[i].height = shared.Cm(1)
table.rows[5].height = shared.Cm(7)
table.rows[6].height = shared.Cm(7)

table.cell(3,1).merge(table.cell(3,3))
table.cell(4,1).merge(table.cell(4,3))
table.cell(5,0).merge(table.cell(5,3))
table.cell(6,0).merge(table.cell(6,3))

table.cell(0,0).paragraphs[0].add_run("학교")
table.cell(0,1).paragraphs[0].add_run("한국디지털미디어고등학교")
table.cell(0,2).paragraphs[0].add_run("전공")
table.cell(0,3).paragraphs[0].add_run(department)
table.cell(1,0).paragraphs[0].add_run("생년월일")
table.cell(1,1).paragraphs[0].add_run(birthday)
table.cell(1,2).paragraphs[0].add_run("학번")
table.cell(1,3).paragraphs[0].add_run(number)
table.cell(2,0).paragraphs[0].add_run("성명")
table.cell(2,1).paragraphs[0].add_run(name)
table.cell(2,2).paragraphs[0].add_run("연락처")
table.cell(2,3).paragraphs[0].add_run(phone)
table.cell(3,0).paragraphs[0].add_run("주소")
table.cell(3,1).paragraphs[0].add_run(address)
table.cell(4,0).paragraphs[0].add_run("자퇴사유")
table.cell(4,1).paragraphs[0].add_run(reason)

for i in range(4):
    for j in range(4):
        table.cell(i,j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

table.cell(5,0).alignment = WD_ALIGN_PARAGRAPH.CENTER
content = table.cell(5,0).paragraphs[0].add_run(f"위와 같은 사유로 한국지털미디어고등학교 학칙에 의거하여 보호자 연서 하에 자퇴서를 제출하오니 허가하여 주시기 바랍니다\n\n{datetime.today().year}년 {datetime.today().month}월 {datetime.today().day}일\n\n신청인 : {name}  (인)\n보호자 : {parent_name}  (인)\n\n\n")
sign_table = table.cell(5,0).add_table(rows = 2, cols = 5)
sign_table.style = doc.styles['Table Grid']
sign_table.direction = WD_TABLE_DIRECTION.LTR
sign_table.columns[0].width = shared.Cm(0.7)
for i in range(1,5):
    sign_table.columns[i].width = shared.Cm(2)
sign_table.rows[0].height = shared.Cm(0.7)
sign_table.rows[1].height = shared.Cm(1.3)
sign_table.cell(0,0).merge(sign_table.cell(1,0))

sign_table.cell(0,0).paragraphs[0].add_run('\n결\n재')
sign_table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
sign_table.cell(0,1).paragraphs[0].add_run('담임')
sign_table.cell(0,2).paragraphs[0].add_run('교무부장')
sign_table.cell(0,3).paragraphs[0].add_run('교감')
sign_table.cell(0,4).paragraphs[0].add_run('교장')

table.cell(6,0).paragraphs[0].add_run("\n\n\n\n\n\n\n\n\n\n\n\n\n한국디지털미디어고등학교장 귀하")
table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i in range(7):
    for j in range(4):
        table.cell(i,j).paragraphs[0].runs[0].font.name = "굴림체"
        table.cell(i,j).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림체') #본문 글씨체
for i in range(5):
    for j in range(2):
        table.cell(i,j).paragraphs[0].runs[0].font.name = "굴림체"
        table.cell(i,j).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림체') #결제란 글씨체

doc.save("자퇴서.docx")